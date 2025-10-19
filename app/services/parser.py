"""
Document parsing service - extracts text from PDF, DOCX, and TXT files.
"""
from pathlib import Path
from typing import Dict, List, Optional
import pdfplumber
from docx import Document as DocxDocument
from app.core.logging import app_logger as logger
import time


class DocumentParser:
    """Parses documents and extracts structured text."""
    
    @staticmethod
    def parse_pdf(file_path: Path) -> Dict:
        """
        Parse PDF file using pdfplumber.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Dict with parsed content and metadata
        """
        start_time = time.time()
        logger.info(f"Parsing PDF: {file_path}")
        
        try:
            pages_data = []
            total_text = ""
            total_words = 0
            has_tables = False
            
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    # Extract text
                    text = page.extract_text() or ""
                    
                    # Extract tables
                    tables = page.extract_tables()
                    page_has_tables = len(tables) > 0 if tables else False
                    if page_has_tables:
                        has_tables = True
                    
                    # Convert tables to text
                    table_text = ""
                    if tables:
                        for table in tables:
                            table_text += "\n[TABLE]\n"
                            for row in table:
                                if row:
                                    table_text += " | ".join(str(cell) if cell else "" for cell in row) + "\n"
                            table_text += "[/TABLE]\n"
                    
                    # Combine text and tables
                    page_content = text + "\n" + table_text if table_text else text
                    
                    # Page metadata
                    page_data = {
                        "page_num": page_num,
                        "text": page_content,
                        "has_tables": page_has_tables,
                        "word_count": len(page_content.split())
                    }
                    
                    pages_data.append(page_data)
                    total_text += page_content + "\n\n"
                    total_words += page_data["word_count"]
                
                processing_time = time.time() - start_time
                
                result = {
                    "raw_text": total_text,
                    "pages": pages_data,
                    "metadata": {
                        "total_pages": len(pdf.pages),
                        "total_words": total_words,
                        "total_chars": len(total_text),
                        "has_tables": has_tables,
                        "processing_time": round(processing_time, 2)
                    }
                }
                
                logger.info(f"PDF parsed successfully: {len(pdf.pages)} pages, {total_words} words in {processing_time:.2f}s")
                return result
                
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {str(e)}")
            raise ValueError(f"Failed to parse PDF: {str(e)}")
    
    @staticmethod
    def parse_docx(file_path: Path) -> Dict:
        """
        Parse DOCX file using python-docx.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dict with parsed content and metadata
        """
        start_time = time.time()
        logger.info(f"Parsing DOCX: {file_path}")
        
        try:
            doc = DocxDocument(file_path)
            
            sections = []
            total_text = ""
            total_words = 0
            has_tables = len(doc.tables) > 0
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    section = {
                        "type": "paragraph",
                        "text": para.text,
                        "style": para.style.name if para.style else "Normal"
                    }
                    sections.append(section)
                    total_text += para.text + "\n"
                    total_words += len(para.text.split())
            
            # Extract tables
            for table_num, table in enumerate(doc.tables, start=1):
                table_text = f"\n[TABLE {table_num}]\n"
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    table_text += row_text + "\n"
                table_text += f"[/TABLE {table_num}]\n"
                
                sections.append({
                    "type": "table",
                    "text": table_text
                })
                total_text += table_text
                total_words += len(table_text.split())
            
            processing_time = time.time() - start_time
            
            result = {
                "raw_text": total_text,
                "sections": sections,
                "metadata": {
                    "total_pages": None,  # DOCX doesn't have page concept
                    "total_words": total_words,
                    "total_chars": len(total_text),
                    "has_tables": has_tables,
                    "num_paragraphs": len([s for s in sections if s["type"] == "paragraph"]),
                    "num_tables": len([s for s in sections if s["type"] == "table"]),
                    "processing_time": round(processing_time, 2)
                }
            }
            
            logger.info(f"DOCX parsed successfully: {total_words} words, {len(doc.tables)} tables in {processing_time:.2f}s")
            return result
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    @staticmethod
    def parse_txt(file_path: Path) -> Dict:
        """
        Parse TXT/MD file.
        
        Args:
            file_path: Path to text file
            
        Returns:
            Dict with parsed content and metadata
        """
        start_time = time.time()
        logger.info(f"Parsing TXT: {file_path}")
        
        try:
            # Try UTF-8 first, fallback to latin-1
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    used_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                raise ValueError("Could not decode file with any supported encoding")
            
            # Split into paragraphs (double newline)
            paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
            
            total_words = len(text.split())
            processing_time = time.time() - start_time
            
            result = {
                "raw_text": text,
                "paragraphs": paragraphs,
                "metadata": {
                    "total_pages": None,
                    "total_words": total_words,
                    "total_chars": len(text),
                    "has_tables": False,
                    "num_paragraphs": len(paragraphs),
                    "encoding": used_encoding,
                    "processing_time": round(processing_time, 2)
                }
            }
            
            logger.info(f"TXT parsed successfully: {total_words} words in {processing_time:.2f}s")
            return result
            
        except Exception as e:
            logger.error(f"Error parsing TXT {file_path}: {str(e)}")
            raise ValueError(f"Failed to parse TXT: {str(e)}")
    
    @classmethod
    def parse_document(cls, file_path: Path, file_type: str) -> Dict:
        """
        Parse document based on file type.
        
        Args:
            file_path: Path to document
            file_type: File extension (.pdf, .docx, .txt)
            
        Returns:
            Dict with parsed content and metadata
        """
        parsers = {
            '.pdf': cls.parse_pdf,
            '.docx': cls.parse_docx,
            '.doc': cls.parse_docx,
            '.txt': cls.parse_txt,
            '.md': cls.parse_txt
        }
        
        parser = parsers.get(file_type.lower())
        if not parser:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        return parser(file_path)
